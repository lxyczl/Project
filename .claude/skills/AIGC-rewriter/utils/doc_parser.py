"""Unified document parser.

Supported formats: .txt / .md / .docx / .pdf (text-based)
Scanned PDFs that yield no text raise ValueError with guidance.
"""

from __future__ import annotations

from pathlib import Path

SUPPORTED_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


def read_document(path_str: str) -> tuple[str, str]:
    """Read document, return (text content, file suffix).

    Suffix returns original suffix (e.g. ".docx"), caller can use it for is_markdown etc.

    Raises:
        FileNotFoundError: File not found
        ValueError: Unsupported format / scanned PDF cannot extract text
    """
    p = Path(path_str)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")

    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"Unsupported format: {suffix}. "
            f"Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    if suffix in (".txt", ".md"):
        text = _read_text_file(p)
    elif suffix == ".docx":
        text = _read_docx(p)
    else:  # .pdf
        text = _read_pdf(p)

    return text, suffix


def _read_text_file(p: Path) -> str:
    """Read plain text file, UTF-8 first, latin-1 fallback."""
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return p.read_text(encoding="latin-1")


def _read_docx(p: Path) -> str:
    """Extract body text from .docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "Requires python-docx: pip install python-docx"
        )

    try:
        doc = Document(str(p))
    except Exception as e:
        if "BadZipFile" in type(e).__name__ or "zipfile" in str(e).lower():
            raise ValueError(
                f"Document {p.name} appears to be corrupted (invalid zip structure). "
                f"Try re-saving the file or converting to .txt/.md format."
            )
        raise

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    if not paragraphs:
        raise ValueError(f"Document {p.name} has no body text content")

    return "\n\n".join(paragraphs)


def _read_pdf(p: Path) -> str:
    """Extract text from text-based PDF.

    Scanned PDFs with no extractable text raise ValueError.
    Encrypted PDFs raise ValueError with clear message.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "Requires PyMuPDF: pip install PyMuPDF"
        )

    doc = fitz.open(str(p))

    # Check for encrypted PDF
    if doc.is_encrypted:
        doc.close()
        raise ValueError(
            f"PDF {p.name} is password-protected. "
            f"Please provide the password or use an unprotected version."
        )

    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()

    if not pages:
        raise ValueError(
            f"PDF {p.name} yielded no text, likely a scanned image PDF. "
            f"Suggestion: Use an OCR tool to convert to text, or copy text into a .txt file."
        )

    return "\n\n".join(pages)


def parse_document(filepath: str | Path) -> list[dict]:
    """Parse document into structured paragraphs.

    Returns list of dicts with:
        - text: paragraph text
        - source_page: page number (PDF only, else None)
        - style: paragraph style (DOCX only, else None)

    Unified output format for all file types.
    """
    p = Path(filepath)
    suffix = p.suffix.lower()

    if suffix in (".txt", ".md"):
        text = _read_text_file(p)
        return [{"text": para.strip(), "source_page": None, "style": None}
                for para in text.split("\n\n") if para.strip()]

    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Requires python-docx: pip install python-docx")

        doc = Document(str(p))
        result = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else None
                result.append({"text": text, "source_page": None, "style": style_name})
        return result

    elif suffix == ".pdf":
        try:
            import fitz
        except ImportError:
            raise ImportError("Requires PyMuPDF: pip install PyMuPDF")

        doc = fitz.open(str(p))
        result = []
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                # Split page text into paragraphs
                for para in text.split("\n\n"):
                    para = para.strip()
                    if para:
                        result.append({"text": para, "source_page": i + 1, "style": None})
        doc.close()
        return result

    else:
        raise ValueError(f"Unsupported format: {suffix}")
