"""
文档解析脚本
从 docx/pdf 文件中提取文本
"""

from pathlib import Path
import re


def parse_docx(file_path: str) -> dict:
    """
    解析 docx 文件

    参数:
        file_path: docx 文件路径

    返回:
        解析结果
    """
    try:
        import docx
    except ImportError:
        return {"error": "需要安装 python-docx: pip install python-docx"}

    path = Path(file_path)
    if not path.exists():
        return {"error": f"文件不存在: {file_path}"}

    try:
        doc = docx.Document(path)

        # 提取段落
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })

        # 识别章节
        sections = identify_sections(paragraphs)

        # 提取元数据
        metadata = extract_metadata(doc)

        return {
            "success": True,
            "file": str(path),
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "sections": sections,
            "metadata": metadata
        }

    except Exception as e:
        return {"error": f"解析失败: {str(e)}"}


def identify_sections(paragraphs: list) -> dict:
    """
    识别论文章节（支持多种格式：带编号/不带编号/全大写/混合大小写）

    参数:
        paragraphs: 段落列表

    返回:
        章节字典
    """
    sections = {
        "title": None,
        "abstract": None,
        "keywords": None,
        "introduction": [],
        "methods": [],
        "results": [],
        "discussion": [],
        "conclusion": [],
        "references": []
    }

    # 章节标题匹配模式（支持编号/无编号/大小写，要求整行匹配）
    section_patterns = [
        (r'^(?:\d+\.?\s*)?introduction\s*$', "introduction"),
        (r'^(?:\d+\.?\s*)?(?:materials?\s+and\s+)?methods?\s*$', "methods"),
        (r'^(?:\d+\.?\s*)?methodology\s*$', "methods"),
        (r'^(?:\d+\.?\s*)?(?:study\s+area|data\s+(?:and\s+)?(?:sources?|collection))\s*$', "methods"),
        (r'^(?:\d+\.?\s*)?results?\s*$', "results"),
        (r'^(?:\d+\.?\s*)?discussion\s*$', "discussion"),
        (r'^(?:\d+\.?\s*)?conclusions?\s*$', "conclusion"),
        (r'^(?:\d+\.?\s*)?(?:summary\s+and\s+)?conclusions?\s*$', "conclusion"),
        (r'^references?\s*$', "references"),
    ]

    current_section = None

    for para in paragraphs:
        text = para["text"].strip()
        text_lower = text.lower()

        # 识别标题（第一个非空段落）
        if para["index"] == 0 and not sections["title"]:
            sections["title"] = text
            continue

        # 识别摘要
        if re.match(r'^(?:\d+\.?\s*)?abstract\b', text_lower) or text_lower == "摘要":
            current_section = "abstract"
            # 提取摘要内容（去掉 "Abstract:" 前缀）
            abstract_content = re.sub(r'^(?:\d+\.?\s*)?abstract\s*[:：]?\s*', '', text, flags=re.IGNORECASE).strip()
            if len(abstract_content) > 50:
                sections["abstract"] = abstract_content
            continue

        # 识别关键词
        if re.match(r'^(?:\d+\.?\s*)?keywords?\b', text_lower) or text_lower.startswith("关键词"):
            current_section = "keywords"
            sections["keywords"] = text
            continue

        # 识别参考文献（优先于其他章节，避免 "References" 被误匹配）
        if re.match(r'^references?\b', text_lower):
            current_section = "references"
            continue

        # 识别其他章节
        matched = False
        for pattern, section_name in section_patterns:
            if re.match(pattern, text_lower):
                current_section = section_name
                matched = True
                break

        if matched:
            continue

        # 添加到当前章节
        if current_section and current_section in sections:
            if isinstance(sections[current_section], list):
                sections[current_section].append(text)
            elif sections[current_section] is None:
                sections[current_section] = text

    # 处理摘要（可能是多段，或嵌入在 "Abstract:" 段落中）
    if sections["abstract"] is None:
        for para in paragraphs[:20]:
            text_lower = para["text"].lower()
            if "abstract" in text_lower and len(para["text"]) > 100:
                idx = text_lower.find("abstract")
                abstract_text = para["text"][idx + len("abstract"):].strip(" :：")
                if len(abstract_text) > 50:
                    sections["abstract"] = abstract_text
                    break

    return sections


def extract_metadata(doc) -> dict:
    """
    提取文档元数据

    参数:
        doc: docx 文档对象

    返回:
        元数据字典
    """
    metadata = {}

    # 核心属性
    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author
    if core.created:
        metadata["created"] = str(core.created)
    if core.modified:
        metadata["modified"] = str(core.modified)

    # 统计信息
    metadata["paragraph_count"] = len(doc.paragraphs)
    metadata["table_count"] = len(doc.tables)

    return metadata


def get_text_for_rewrite(parse_result: dict, section: str = None) -> str:
    """
    获取用于改写的文本

    参数:
        parse_result: 解析结果
        section: 指定章节 (可选)

    返回:
        文本字符串
    """
    if "error" in parse_result:
        return ""

    if section:
        # 获取指定章节
        section_lower = section.lower()
        if section_lower in parse_result["sections"]:
            content = parse_result["sections"][section_lower]
            if isinstance(content, list):
                return "\n\n".join(content)
            elif content:
                return content
        return ""

    # 获取所有文本
    texts = [p["text"] for p in parse_result["paragraphs"]]
    return "\n\n".join(texts)


def parse_pdf(file_path: str) -> dict:
    """
    解析 PDF 文件 (简单版本)

    参数:
        file_path: PDF 文件路径

    返回:
        解析结果
    """
    try:
        import PyPDF2
    except ImportError:
        return {"error": "需要安装 PyPDF2: pip install PyPDF2"}

    path = Path(file_path)
    if not path.exists():
        return {"error": f"文件不存在: {file_path}"}

    try:
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            paragraphs = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # 按段落分割
                    for j, para_text in enumerate(text.split('\n\n')):
                        para_text = para_text.strip()
                        if para_text and len(para_text) > 10:
                            paragraphs.append({
                                "index": len(paragraphs),
                                "text": para_text,
                                "page": i + 1,
                                "style": "Normal"
                            })

            # 识别章节
            sections = identify_sections(paragraphs)

            return {
                "success": True,
                "file": str(path),
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs,
                "sections": sections,
                "metadata": {
                    "page_count": len(reader.pages)
                }
            }

    except Exception as e:
        return {"error": f"解析失败: {str(e)}"}


def parse_document(file_path: str) -> dict:
    """
    解析文档 (自动识别格式)

    参数:
        file_path: 文件路径

    返回:
        解析结果
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == '.docx':
        return parse_docx(file_path)
    elif suffix == '.pdf':
        return parse_pdf(file_path)
    else:
        return {"error": f"不支持的文件格式: {suffix}"}
